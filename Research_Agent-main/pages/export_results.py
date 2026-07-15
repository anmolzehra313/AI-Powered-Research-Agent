"""
Export Results Page

Central export page for all session data.
Supports export in CSV, PDF, DOCX, and BibTeX formats.
"""

import streamlit as st
import pandas as pd
import io
import sys
import os
from datetime import datetime

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))


def render():
    """Render the Export Results page."""
    st.title("📥 Export Results & Reports")
    st.markdown("Export your research data in multiple formats: **CSV**, **PDF**, **DOCX**.")

    # Check what data is available
    available_data = _check_available_data()

    if not any(available_data.values()):
        st.info("No data available for export yet. Use the other pages to generate results first.")
        st.markdown("""
        **Available export sources:**
        - 🔍 **Search Results** – Search for papers on the Search Papers page
        - 📚 **Citations** – Extract citations on the Citation Extractor page
        - 🔎 **Plagiarism Report** – Run plagiarism check on the Plagiarism Checker page
        - 🤖 **AI Chat History** – Chat with papers on the AI Chat page
        - 📝 **Formatted References** – Format references on the Reference Formatter page
        """)
        return

    st.markdown("---")

    # 1. Search Results Export
    if available_data['search_results']:
        st.subheader("🔍 Search Results")
        papers = st.session_state['search_results']
        st.info(f"{len(papers)} papers available for export")

        col1, col2, col3 = st.columns(3)

        with col1:
            csv_data = _papers_to_csv(papers)
            st.download_button(
                "📥 CSV", data=csv_data,
                file_name="search_results.csv", mime="text/csv",
                use_container_width=True,
                key="download_search_csv"
            )

        with col2:
            pdf_data = _generate_pdf_report("Search Results", _papers_to_text(papers))
            if pdf_data:
                st.download_button(
                    "📥 PDF", data=pdf_data,
                    file_name="search_results.pdf", mime="application/pdf",
                    use_container_width=True,
                    key="download_search_pdf"
                )

        with col3:
            docx_data = _generate_docx_report("Search Results", _papers_to_text(papers))
            if docx_data:
                st.download_button(
                    "📥 DOCX", data=docx_data,
                    file_name="search_results.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="download_search_docx"
                )

        st.markdown("---")

    # 2. Citations Export
    if available_data['citations']:
        st.subheader("📚 Extracted Citations")
        citations = st.session_state['extracted_citations']
        st.info(f"{len(citations)} citations available for export")

        col1, col2, col3 = st.columns(3)

        with col1:
            csv_data = st.session_state.get('citations_csv', _citations_to_csv(citations))
            st.download_button(
                "📥 CSV", data=csv_data,
                file_name="citations.csv", mime="text/csv",
                use_container_width=True,
                key="download_citations_csv"
            )

        with col2:
            pdf_data = _generate_pdf_report("Extracted Citations", _citations_to_text(citations))
            if pdf_data:
                st.download_button(
                    "📥 PDF", data=pdf_data,
                    file_name="citations.pdf", mime="application/pdf",
                    use_container_width=True,
                    key="download_citations_pdf"
                )

        with col3:
            docx_data = _generate_docx_report("Extracted Citations", _citations_to_text(citations))
            if docx_data:
                st.download_button(
                    "📥 DOCX", data=docx_data,
                    file_name="citations.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="download_citations_docx"
                )

        st.markdown("---")

    # 3. Plagiarism Report Export
    if available_data['plagiarism']:
        st.subheader("🔎 Plagiarism Report")
        report = st.session_state.get('plagiarism_report', '')
        result = st.session_state.get('plagiarism_result', {})
        st.info(f"Overall score: {result.get('overall_plagiarism_score', 'N/A')}%")

        col1, col2, col3 = st.columns(3)

        with col1:
            st.download_button(
                "📥 TXT", data=report,
                file_name="plagiarism_report.txt", mime="text/plain",
                use_container_width=True,
                key="download_plag_txt"
            )

        with col2:
            pdf_data = _generate_pdf_report("Plagiarism Report", report)
            if pdf_data:
                st.download_button(
                    "📥 PDF", data=pdf_data,
                    file_name="plagiarism_report.pdf", mime="application/pdf",
                    use_container_width=True,
                    key="download_plag_pdf"
                )

        with col3:
            docx_data = _generate_docx_report("Plagiarism Report", report)
            if docx_data:
                st.download_button(
                    "📥 DOCX", data=docx_data,
                    file_name="plagiarism_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="download_plag_docx"
                )

        st.markdown("---")

    # 4. AI Chat History Export
    if available_data['chat_history']:
        st.subheader("🤖 AI Chat History")
        messages = st.session_state['chat_messages']
        st.info(f"{len(messages)} messages in chat history")

        chat_text = _chat_to_text(messages)

        col1, col2, col3 = st.columns(3)

        with col1:
            st.download_button(
                "📥 TXT", data=chat_text,
                file_name="ai_chat_history.txt", mime="text/plain",
                use_container_width=True,
                key="download_chat_txt"
            )

        with col2:
            pdf_data = _generate_pdf_report("AI Chat History", chat_text)
            if pdf_data:
                st.download_button(
                    "📥 PDF", data=pdf_data,
                    file_name="ai_chat_history.pdf", mime="application/pdf",
                    use_container_width=True,
                    key="download_chat_pdf"
                )

        with col3:
            docx_data = _generate_docx_report("AI Chat History", chat_text)
            if docx_data:
                st.download_button(
                    "📥 DOCX", data=docx_data,
                    file_name="ai_chat_history.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    key="download_chat_docx"
                )

        st.markdown("---")

    # 5. Formatted References Export
    if available_data['formatted_refs']:
        st.subheader("📝 Formatted References")

        if 'batch_formatted' in st.session_state:
            st.download_button(
                "📥 Formatted References (TXT)",
                data=st.session_state['batch_formatted'],
                file_name="formatted_references.txt", mime="text/plain",
                use_container_width=True,
            )

    # Export All Button
    st.markdown("---")
    st.subheader("📦 Export All Data")
    
    # Generate combined report if any data exists
    if any(available_data.values()):
        combined = _generate_combined_report(available_data)
        if combined:
            st.download_button(
                "📥 Download Combined Report (TXT)",
                data=combined,
                file_name="research_agent_report.txt",
                mime="text/plain",
                use_container_width=True,
                key="download_combined_report"
            )
    else:
        st.info("Generate some research data to enable the combined report.")


# ==================== Helper Functions ====================

def _check_available_data():
    """Check what data is available in session state."""
    return {
        'search_results': 'search_results' in st.session_state and bool(st.session_state['search_results']),
        'citations': 'extracted_citations' in st.session_state and bool(st.session_state['extracted_citations']),
        'plagiarism': 'plagiarism_report' in st.session_state and bool(st.session_state['plagiarism_report']),
        'chat_history': 'chat_messages' in st.session_state and bool(st.session_state['chat_messages']),
        'formatted_refs': ('batch_formatted' in st.session_state or 'batch_bibtex' in st.session_state),
    }


def _papers_to_csv(papers):
    """Convert papers list to CSV bytes with UTF-8-SIG."""
    df = pd.DataFrame(papers)
    return df.to_csv(index=False).encode('utf-8-sig')


def _papers_to_text(papers):
    """Convert papers list to formatted text."""
    lines = []
    for i, p in enumerate(papers, 1):
        lines.append(f"--- Paper {i} ---")
        lines.append(f"Title: {p.get('title', 'N/A')}")
        lines.append(f"Authors: {p.get('authors', 'N/A')}")
        lines.append(f"Year: {p.get('year', 'N/A')}")
        lines.append(f"Citations: {p.get('citation_count', 'N/A')}")
        lines.append(f"DOI: {p.get('doi', 'N/A')}")
        lines.append(f"Source: {p.get('source', 'N/A')}")
        lines.append(f"Abstract: {p.get('abstract', 'N/A')}")
        lines.append("")
    return '\n'.join(lines)


def _citations_to_csv(citations):
    """Convert citations to CSV bytes with UTF-8-SIG."""
    df = pd.DataFrame(citations)
    return df.to_csv(index=False).encode('utf-8-sig')


def _citations_to_text(citations):
    """Convert citations to formatted text."""
    lines = []
    for i, c in enumerate(citations, 1):
        lines.append(f"[{i}] {c.get('authors', '')} ({c.get('year', '')}). "
                     f"{c.get('title', '')}. {c.get('journal', '')}.")
        if c.get('doi'):
            lines.append(f"    DOI: {c['doi']}")
        lines.append("")
    return '\n'.join(lines)


def _chat_to_text(messages):
    """Convert chat messages to text."""
    lines = [f"AI Research Assistant - Chat History", f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ""]
    for msg in messages:
        role = "You" if msg['role'] == 'user' else "AI Assistant"
        lines.append(f"{role}: {msg['content']}")
        lines.append("")
    return '\n'.join(lines)


def _generate_pdf_report(title, content):
    """Generate a PDF report using FPDF."""
    try:
        from fpdf import FPDF

        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)

        # Title
        pdf.set_font('Helvetica', 'B', 16)
        pdf.cell(0, 10, title, ln=True, align='C')
        pdf.ln(5)

        # Date
        pdf.set_font('Helvetica', 'I', 10)
        pdf.cell(0, 10, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ln=True, align='C')
        pdf.ln(10)

        # Content
        pdf.set_font('Helvetica', '', 10)
        # Handle encoding - replace problematic characters
        safe_content = content.encode('latin-1', errors='replace').decode('latin-1')
        pdf.multi_cell(0, 5, safe_content)

        # Convert bytearray to bytes for Streamlit compatibility
        return bytes(pdf.output())

    except Exception as e:
        st.warning(f"PDF generation failed: {e}")
        return None


def _generate_docx_report(title, content):
    """Generate a DOCX report using python-docx."""
    try:
        
        from docx import Document

        doc = Document()
        doc.add_heading(title, 0)
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph("")

        # Add content paragraphs
        for line in content.split('\n'):
            if line.strip():
                doc.add_paragraph(line)

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    except Exception as e:
        st.warning(f"DOCX generation failed: {e}")
        return None


def _generate_combined_report(available_data):
    """Generate a combined text report of all available data."""
    report = []
    report.append("=" * 60)
    report.append("RESEARCH AGENT - COMBINED REPORT")
    report.append(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    report.append("=" * 60)
    report.append("")

    if available_data['search_results']:
        papers = st.session_state['search_results']
        report.append("## SEARCH RESULTS")
        report.append(_papers_to_text(papers))
        report.append("")

    if available_data['citations']:
        citations = st.session_state['extracted_citations']
        report.append("## EXTRACTED CITATIONS")
        report.append(_citations_to_text(citations))
        report.append("")

    if available_data['plagiarism']:
        report.append("## PLAGIARISM REPORT")
        report.append(st.session_state.get('plagiarism_report', ''))
        report.append("")

    if available_data['chat_history']:
        messages = st.session_state['chat_messages']
        report.append("## AI CHAT HISTORY")
        report.append(_chat_to_text(messages))
        report.append("")

    if available_data['formatted_refs']:
        report.append("## FORMATTED REFERENCES")
        report.append(st.session_state.get('batch_formatted', ''))
        report.append("")

    return '\n'.join(report)
